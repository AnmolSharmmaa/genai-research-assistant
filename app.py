import streamlit as st
from backend.file_loader import load_document
from backend.summarizer import summarize_text
from backend.QA import answer_question
from backend.challenge import generate_questions, evaluate_answer
from io import BytesIO
from docx import Document
import speech_recognition as sr
from gtts import gTTS
import os
from tempfile import NamedTemporaryFile

def record_voice():
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        st.info("🎤 Listening... Please speak your question.")
        audio = recognizer.listen(source, timeout=5)
    try:
        text = recognizer.recognize_google(audio)
        st.success(f"🗣️ You said: {text}")
        return text
    except sr.UnknownValueError:
        st.error("😕 Sorry, I couldn't understand.")
        return None
    except sr.RequestError:
        st.error("🚫 Speech recognition service error.")
        return None

# Page configuration
st.set_page_config(page_title="GenAI Research Assistant", layout="wide", initial_sidebar_state="collapsed")

if "qna_pairs" not in st.session_state:
    st.session_state.qna_pairs = []

if "challenge_feedback" not in st.session_state:
    st.session_state.challenge_feedback = []

# --- CUSTOM CSS STYLING ---
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Roboto:wght@300;400;700&display=swap');

    html, body, [class*="st-"] {
        font-family: 'Roboto', sans-serif;
    }

    /* Main App Background */
    [data-testid="stAppViewContainer"] {
        background-color: #0d1117;
        background-image: linear-gradient(160deg, #0d1117 0%, #212834 100%);
        color: #e6edf3;
    }

    /* --- NEW ANIMATED SEPARATOR --- */
    @keyframes flow-animation {
        0% { background-position: 200% 0; }
        100% { background-position: -200% 0; }
    }

    .animated-separator {
        height: 3px;
        background: linear-gradient(90deg, transparent, #00c6ff, transparent);
        background-size: 45% 100%;
        animation: flow-animation 4s linear infinite;
        border-radius: 5px;
        margin: 1.5rem 0;
    }


    /* Main header styling */
    h1, h2 {
        color: #ffffff;
        text-shadow: 2px 2px 5px rgba(0, 191, 255, 0.5);
    }

    h3, h4, h5 {
        color: #c9d1d9;
    }

    /* Card-like containers for content */
    .card {
        background-color: rgba(22, 27, 34, 0.7);
        border-radius: 12px;
        padding: 2em;
        margin-bottom: 1em;
        border: 1px solid rgba(255, 255, 255, 0.1);
        box-shadow: 0 4px 30px rgba(0, 0, 0, 0.1);
        backdrop-filter: blur(10px);
        -webkit-backdrop-filter: blur(10px);
        transition: all 0.3s ease;
    }

    .card:hover {
        border: 1px solid rgba(0, 191, 255, 0.5);
        box-shadow: 0 4px 30px rgba(0, 191, 255, 0.2);
    }

    /* Welcome Screen Styles */
    .welcome-container {
        padding: 2rem;
        text-align: center;
    }

    .welcome-title {
        font-size: 2.5rem;
        font-weight: 700;
        background: -webkit-linear-gradient(45deg, #00c6ff, #0072ff);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        padding-bottom: 0.5rem;
    }

    .feature-box {
        padding: 1.5rem;
        border-radius: 10px;
        background-color: rgba(13, 17, 23, 0.5);
        transition: all 0.3s ease-in-out;
        height: 100%;
    }

    .feature-box:hover {
        transform: translateY(-5px);
        background-color: rgba(13, 17, 23, 0.8);
    }

    .feature-icon {
        font-size: 3rem;
        animation: float 3s ease-in-out infinite;
    }

    @keyframes float {
        0% { transform: translateY(0px); }
        50% { transform: translateY(-8px); }
        100% { transform: translateY(0px); }
    }

    .icon-1 { animation-delay: 0s; }
    .icon-2 { animation-delay: 0.5s; }
    .icon-3 { animation-delay: 1s; }

    section[data-testid="stFileUploader"] > div {
        border-color: rgba(0, 191, 255, 0.4);
        background-color: rgba(13, 17, 23, 0.5);
    }

    section[data-testid="stFileUploader"] label {
        color: #e6edf3;
        font-weight: 500;
    }

    .stButton>button {
        background-image: linear-gradient(to right, #00c6ff 0%, #0072ff 51%, #00c6ff 100%);
        color: white;
        padding: 0.6em 1.5em;
        text-align: center;
        text-transform: uppercase;
        transition: 0.5s;
        background-size: 200% auto;
        border: none;
        border-radius: 8px;
        display: block;
        font-weight: 700;
    }

    .stButton>button:hover {
        background-position: right center;
        color: #fff;
        text-decoration: none;
        border: none;
    }

    div[data-testid="stTextInput"] input, div[data-testid="stTextArea"] textarea {
        background-color: rgba(13, 17, 23, 0.8);
        color: #e6edf3 !important;
        border: 1px solid rgba(255, 255, 255, 0.1);
        border-radius: 8px;
        padding: 0.75em;
    }

    div[data-testid="stRadio"] > label {
        background-color: rgba(33, 40, 52, 0.7);
        padding: 0.75em 1em;
        border-radius: 8px;
        margin-bottom: 0.5em;
        border: 1px solid rgba(255, 255, 255, 0.1);
        transition: all 0.2s ease-in-out;
    }

    div[data-testid="stRadio"] > label:hover {
        background-color: rgba(0, 191, 255, 0.2);
        border-color: rgba(0, 191, 255, 0.7);
    }

    div[data-testid="stAlert-success"], div[data-testid="stInfo"] {
        background-color: rgba(33, 40, 52, 0.7);
        border: 1px solid #00c6ff;
        border-radius: 8px;
        color: #e6edf3;
    }
    </style>
""", unsafe_allow_html=True)


# --- APP LAYOUT ---

st.markdown("<h1 style='text-align: center;'>✨ Smart Research Assistant ✨</h1>", unsafe_allow_html=True)

st.markdown("<div class='animated-separator'></div>", unsafe_allow_html=True)

st.markdown("<div class='card'>", unsafe_allow_html=True)

st.markdown("""
<h3 style="
    text-align: center;
    font-weight: 700;
    background: -webkit-linear-gradient(45deg, #00c6ff, #0072ff);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
">
STEP 1: Upload Your Document
</h3>
""", unsafe_allow_html=True)

uploaded_file = st.file_uploader(
    "Upload a PDF or TXT file to begin your intelligent analysis.",
    type=["pdf", "txt"],
    label_visibility="visible"
)
st.markdown("</div>", unsafe_allow_html=True)


# --- CONDITIONAL DISPLAY LOGIC ---

if not uploaded_file:
    st.markdown("<div class='card welcome-container'>", unsafe_allow_html=True)
    st.markdown("<p class='welcome-title'>Unlock the Power of Your Documents</p>", unsafe_allow_html=True)
    st.write("Your journey into intelligent document analysis starts here. Simply upload a file to begin.")
    
    col1, col2, col3 = st.columns(3, gap="large")
    with col1:
        st.markdown("""
            <div class='feature-box'>
                <div class='feature-icon icon-1'>📤</div>
                <h4>Instant Summaries</h4>
                <p>Get the gist of any document in seconds with AI-powered summarization.</p>
            </div>
        """, unsafe_allow_html=True)
    with col2:
        st.markdown("""
            <div class='feature-box'>
                <div class='feature-icon icon-2'>💬</div>
                <h4>Dynamic Q&A</h4>
                <p>Ask direct questions and receive precise answers sourced directly from your text.</p>
            </div>
        """, unsafe_allow_html=True)
    with col3:
        st.markdown("""
            <div class='feature-box'>
                <div class='feature-icon icon-3'>🧠</div>
                <h4>Logic Challenges</h4>
                <p>Test your comprehension with auto-generated logical questions based on the content.</p>
            </div>
        """, unsafe_allow_html=True)
        
    st.markdown("</div>", unsafe_allow_html=True)

if uploaded_file:
    st.success(f"✅ Success! You've uploaded **{uploaded_file.name}**.")

    raw_text = load_document(uploaded_file)

    st.subheader("📄 Auto Summary")
    with st.spinner("🔄 Generating summary..."):
        summary = summarize_text(raw_text)
    st.info(summary)

    st.markdown("<div class='card'>", unsafe_allow_html=True)
    st.subheader("🧠 STEP 2: Interact with the Document")

    mode = st.radio(
        "Choose interaction mode:",
        ["❓ Ask Anything", "🧠 Challenge Me"]
    )

    if mode == "❓ Ask Anything":
        st.markdown("<h5>🔍 Ask a question about the document</h5>", unsafe_allow_html=True)

        question = st.text_input("Type your question here:")
        if question:
            with st.spinner("🤖 Generating answer..."):
                answer, justification = answer_question(question, raw_text)
            st.session_state.qna_pairs.append((question, answer, justification))
            st.success(f"**Answer:** {answer}")
            st.info(f"**Justification:** {justification}")

        if st.button("🎙️ Ask by Voice"):
            voice_input = record_voice()
            if voice_input:
                with st.spinner("🤖 Generating answer..."):
                    answer, justification = answer_question(voice_input, raw_text)
                st.session_state.qna_pairs.append((voice_input, answer, justification))
                st.markdown(f"🧠 **Answer:** {answer}")
                st.markdown(
                    f"📌 **Supporting Snippet:**<br>"
                    f"<span style='background-color:#003366;padding:4px;border-radius:6px;color:#ffffff'>{justification}</span>",
                    unsafe_allow_html=True
                )

                tts = gTTS(text=answer, lang='en')
                with NamedTemporaryFile(delete=False, suffix=".mp3") as temp_file:
                    tts.save(temp_file.name)
                    st.audio(temp_file.name, format="audio/mp3")

    elif mode == "🧠 Challenge Me":
        st.markdown("<h5>🧠 Test your understanding with logic-based questions:</h5>", unsafe_allow_html=True)
        questions = generate_questions(raw_text)
        for i, q in enumerate(questions):
            st.markdown(f"**Q{i+1}:** {q['question']}")
            user_ans = st.text_input(f"Your answer to Q{i+1}:", key=f"challenge_q{i}")
            if user_ans:
                correct, feedback = evaluate_answer(user_ans, q['answer'])
                if correct:
                    st.success(f"✅ {feedback}")
                else:
                    st.error(f"❌ {feedback}")
                    st.info(f"📌 Justification: {q.get('justification', 'Based on document.')}")

    st.markdown("</div>", unsafe_allow_html=True)

    # --- DOWNLOAD REPORT ---
    def generate_docx_report(summary, qna_pairs=[], challenge_feedback=[]):
        doc = Document()
        doc.add_heading("GenAI Research Assistant Report", 0)

        doc.add_heading("Auto Summary", level=1)
        doc.add_paragraph(summary)

        if qna_pairs:
            doc.add_heading("Ask Anything - Q&A", level=1)
            for q, a, j in qna_pairs:
                doc.add_paragraph(f"Q: {q}", style='List Bullet')
                doc.add_paragraph(f"A: {a}")
                doc.add_paragraph(f"Justification: {j}")

        if challenge_feedback:
            doc.add_heading("Challenge Me - Results", level=1)
            for i, (question, user_ans, feedback) in enumerate(challenge_feedback, 1):
                doc.add_paragraph(f"Q{i}: {question}", style='List Bullet')
                doc.add_paragraph(f"Your Answer: {user_ans}")
                doc.add_paragraph(f"Feedback: {feedback}")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    st.markdown("<div class='card'>", unsafe_allow_html=True)
    st.subheader("📥 Download Your Report")

    docx_data = generate_docx_report(
        summary=summary,
        qna_pairs=st.session_state.qna_pairs,
        challenge_feedback=st.session_state.challenge_feedback
    )

    st.download_button(
        label="📄 Download Report (.docx)",
        data=docx_data,
        file_name="GenAI_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    st.markdown("</div>", unsafe_allow_html=True)
